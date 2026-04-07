from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.core.config import get_settings


@dataclass
class GeneratedOutputs:
    docx_bytes: bytes | None = None
    pdf_html: str | None = None  # Print-ready HTML (user can print-to-PDF)
    email_confident: str = ""
    email_warm: str = ""


class DocumentGenerator:
    """Generates DOCX, print-ready HTML, and email drafts from proposal data."""

    def generate_all(
        self,
        proposal: dict,
        agency_name: str,
        agency_colours: dict | None = None,
    ) -> GeneratedOutputs:
        docx_bytes = self._generate_docx(proposal, agency_name)
        pdf_html = self._generate_pdf_html(proposal, agency_name, agency_colours)
        email_confident, email_warm = self._generate_email_drafts(proposal, agency_name)

        return GeneratedOutputs(
            docx_bytes=docx_bytes,
            pdf_html=pdf_html,
            email_confident=email_confident,
            email_warm=email_warm,
        )

    def _generate_docx(self, proposal: dict, agency_name: str) -> bytes:
        """Generate a professional DOCX proposal document."""
        doc = Document()

        # Styles
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)

        # 1. Cover page
        doc.add_paragraph("")
        doc.add_paragraph("")
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(agency_name)
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(26, 26, 26)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(proposal.get("project_name", "Proposal"))
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(100, 100, 100)

        client_name = proposal.get("brief", {}).get("client", {}).get("name", "")
        if client_name:
            client_p = doc.add_paragraph()
            client_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = client_p.add_run(f"Prepared for {client_name}")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(120, 120, 120)

        doc.add_paragraph("\n\nConfidential")
        doc.add_page_break()

        # 2. Covering Letter
        if proposal.get("covering_letter"):
            doc.add_heading("Covering Letter", level=1)
            for para in proposal["covering_letter"].split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.add_page_break()

        # 3. Executive Summary
        if proposal.get("executive_summary"):
            doc.add_heading("Executive Summary", level=1)
            for para in proposal["executive_summary"].split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.add_page_break()

        # 4. Scope of Engagement
        scope_sections = proposal.get("scope_sections", [])
        if scope_sections:
            doc.add_heading("Scope of Engagement", level=1)
            for section in scope_sections:
                doc.add_heading(section.get("deliverable", "Deliverable"), level=2)
                content = section.get("content", "")
                for para in content.split("\n"):
                    line = para.strip()
                    if line.startswith("###"):
                        doc.add_heading(line.lstrip("# "), level=3)
                    elif line.startswith("**") and line.endswith("**"):
                        p = doc.add_paragraph()
                        run = p.add_run(line.strip("*"))
                        run.bold = True
                    elif line.startswith("- "):
                        doc.add_paragraph(line[2:], style="List Bullet")
                    elif line:
                        doc.add_paragraph(line)
            doc.add_page_break()

        # 5. Cost Summary
        cost_model = proposal.get("cost_model", {})
        if cost_model.get("line_items"):
            doc.add_heading("Investment Summary", level=1)

            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Grid Accent 1"
            headers = table.rows[0].cells
            headers[0].text = "Deliverable"
            headers[1].text = "Quantity"
            headers[2].text = "Unit Cost"
            headers[3].text = "Total"

            for item in cost_model["line_items"]:
                row = table.add_row().cells
                row[0].text = item.get("deliverable", "")
                row[1].text = str(item.get("quantity", 1))
                row[2].text = f"₹{item.get('unit_cost', 0):,}"
                row[3].text = f"₹{item.get('total', 0):,}"

            doc.add_paragraph("")
            doc.add_paragraph(f"Subtotal: ₹{cost_model.get('subtotal', 0):,}")
            if cost_model.get("discount_percent", 0) > 0:
                doc.add_paragraph(f"Discount ({cost_model['discount_percent']}%): -₹{cost_model.get('discount_amount', 0):,}")
            doc.add_paragraph(f"Total (excl. GST): ₹{cost_model.get('total', 0):,}")
            doc.add_paragraph(f"GST (18%): ₹{cost_model.get('gst_amount', 0):,}")
            p = doc.add_paragraph()
            run = p.add_run(f"Grand Total: ₹{cost_model.get('grand_total', 0):,}")
            run.bold = True
            run.font.size = Pt(14)
            doc.add_page_break()

        # 6. Cost Rationale
        if proposal.get("cost_rationale"):
            doc.add_heading("Cost Rationale", level=1)
            for para in proposal["cost_rationale"].split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.add_page_break()

        # 7. Terms & Conditions
        if proposal.get("terms"):
            doc.add_heading("Terms & Conditions", level=1)
            for para in proposal["terms"].split("\n"):
                line = para.strip()
                if line.startswith("###"):
                    doc.add_heading(line.lstrip("# "), level=2)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line:
                    doc.add_paragraph(line)

        # 8. Contact
        doc.add_page_break()
        doc.add_heading("Contact", level=1)
        doc.add_paragraph(agency_name)
        doc.add_paragraph("For questions about this proposal, please contact us directly.")

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _generate_pdf_html(
        self, proposal: dict, agency_name: str, agency_colours: dict | None = None,
    ) -> str:
        """Generate a print-ready HTML document (can be saved as PDF via browser print)."""
        primary = (agency_colours or {}).get("primary", "#1a1a1a")
        client_name = proposal.get("brief", {}).get("client", {}).get("name", "")
        project_name = proposal.get("project_name", "Proposal")
        cost_model = proposal.get("cost_model", {})

        # Build cost table rows
        cost_rows = ""
        for item in cost_model.get("line_items", []):
            cost_rows += f"""<tr>
                <td>{item.get('deliverable','')}</td>
                <td style="text-align:center">{item.get('quantity',1)}</td>
                <td style="text-align:right">₹{item.get('unit_cost',0):,}</td>
                <td style="text-align:right">₹{item.get('total',0):,}</td>
            </tr>"""

        # Scope sections HTML
        scope_html = ""
        for section in proposal.get("scope_sections", []):
            scope_html += f"""
            <div class="scope-section">
                <h3>{section.get('deliverable','')}</h3>
                <div class="scope-content">{self._md_to_html(section.get('content',''))}</div>
            </div>"""

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{project_name} - {agency_name}</title>
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display&family=DM+Sans:wght@400;500;700&display=swap');
@media print {{ @page {{ margin: 2cm; }} }}
body {{ font-family: 'DM Sans', sans-serif; color: #1a1a1a; line-height: 1.7; max-width: 800px; margin: 0 auto; padding: 40px; }}
h1 {{ font-family: 'DM Serif Display', serif; font-size: 28px; color: {primary}; margin-top: 48px; }}
h2 {{ font-family: 'DM Serif Display', serif; font-size: 20px; color: {primary}; margin-top: 32px; }}
h3 {{ font-size: 16px; font-weight: 700; margin-top: 24px; }}
.cover {{ text-align: center; padding: 120px 0 80px; page-break-after: always; }}
.cover h1 {{ font-size: 36px; margin-bottom: 8px; }}
.cover .client {{ font-size: 18px; color: #666; }}
.cover .conf {{ font-size: 12px; color: #999; margin-top: 60px; }}
table {{ width: 100%; border-collapse: collapse; margin: 16px 0; }}
th {{ background: #f5f5f4; text-align: left; padding: 8px 12px; font-size: 12px; text-transform: uppercase; letter-spacing: 0.5px; border-bottom: 2px solid #e5e5e4; }}
td {{ padding: 8px 12px; border-bottom: 1px solid #e5e5e4; font-size: 14px; }}
.totals {{ margin-top: 16px; text-align: right; }}
.totals .grand {{ font-size: 18px; font-weight: 700; }}
.scope-section {{ margin: 24px 0; padding: 16px; background: #fafaf9; border-radius: 8px; }}
.page-break {{ page-break-before: always; }}
</style>
</head>
<body>
<div class="cover">
    <h1>{agency_name}</h1>
    <p style="font-size:22px;color:#666">{project_name}</p>
    {'<p class="client">Prepared for ' + client_name + '</p>' if client_name else ''}
    <p class="conf">Confidential</p>
</div>

<h1>Covering Letter</h1>
{self._md_to_html(proposal.get('covering_letter', ''))}

<div class="page-break"></div>
<h1>Executive Summary</h1>
{self._md_to_html(proposal.get('executive_summary', ''))}

<div class="page-break"></div>
<h1>Scope of Engagement</h1>
{scope_html}

<div class="page-break"></div>
<h1>Investment Summary</h1>
<table>
<thead><tr><th>Deliverable</th><th style="text-align:center">Qty</th><th style="text-align:right">Unit Cost</th><th style="text-align:right">Total</th></tr></thead>
<tbody>{cost_rows}</tbody>
</table>
<div class="totals">
    <p>Subtotal: ₹{cost_model.get('subtotal',0):,}</p>
    {'<p>Discount (' + str(cost_model.get("discount_percent",0)) + '%): -₹' + f'{cost_model.get("discount_amount",0):,}' + '</p>' if cost_model.get('discount_percent',0) > 0 else ''}
    <p>Total (excl. GST): ₹{cost_model.get('total',0):,}</p>
    <p>GST (18%): ₹{cost_model.get('gst_amount',0):,}</p>
    <p class="grand">Grand Total: ₹{cost_model.get('grand_total',0):,}</p>
</div>

{'<div class="page-break"></div><h1>Cost Rationale</h1>' + self._md_to_html(proposal.get('cost_rationale','')) if proposal.get('cost_rationale') else ''}

<div class="page-break"></div>
<h1>Terms &amp; Conditions</h1>
{self._md_to_html(proposal.get('terms', ''))}

<div class="page-break"></div>
<h1>Contact</h1>
<p>{agency_name}</p>
<p>For questions about this proposal, please contact us directly.</p>
</body>
</html>"""

    def _generate_email_drafts(self, proposal: dict, agency_name: str) -> tuple[str, str]:
        """Generate two email draft variants."""
        client_name = proposal.get("brief", {}).get("client", {}).get("name", "the client")
        project_name = proposal.get("project_name", "the project")
        contact_name = ""
        contacts = proposal.get("brief", {}).get("client", {}).get("contacts", [])
        if contacts and isinstance(contacts[0], dict):
            contact_name = contacts[0].get("name", "")

        greeting = f"Hi {contact_name}" if contact_name else "Hi"

        confident = f"""{greeting},

We've put together something for the {project_name}. Rather than a static PDF, we've built an interactive proposal you can explore at your own pace:

[PROPOSAL LINK]

The proposal covers our approach, detailed scope for each deliverable, investment breakdown with market benchmarks, and a suggested timeline.

I'd love to walk you through it. Would Thursday at 3pm work for a 30-minute session?

Best,
Karthik
{agency_name}"""

        warm = f"""{greeting},

Following our conversation about the {project_name}, I've put together a comprehensive proposal covering the scope, approach, and investment.

You can view it here: [PROPOSAL LINK]

I've also attached the PDF version for your records.

Happy to jump on a call this week to discuss any questions. What works best for you?

Warm regards,
Karthik
{agency_name}"""

        return confident, warm

    @staticmethod
    def _md_to_html(text: str) -> str:
        """Very basic markdown to HTML conversion for proposal content."""
        if not text:
            return ""
        lines = text.split("\n")
        html_parts = []
        for line in lines:
            stripped = line.strip()
            if not stripped:
                html_parts.append("")
            elif stripped.startswith("### "):
                html_parts.append(f"<h3>{stripped[4:]}</h3>")
            elif stripped.startswith("## "):
                html_parts.append(f"<h2>{stripped[3:]}</h2>")
            elif stripped.startswith("# "):
                html_parts.append(f"<h2>{stripped[2:]}</h2>")
            elif stripped.startswith("- "):
                html_parts.append(f"<li>{stripped[2:]}</li>")
            elif stripped.startswith("**") and stripped.endswith("**"):
                html_parts.append(f"<p><strong>{stripped[2:-2]}</strong></p>")
            else:
                html_parts.append(f"<p>{stripped}</p>")
        return "\n".join(html_parts)
